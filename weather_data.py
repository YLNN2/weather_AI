import random
from datetime import date, timedelta
from docx import Document

document = Document()
document.add_heading('Погода в городе', 0)

#year = random.randrange(1900, 2021)
year = 2017
winds = ['северный', 'южный', 'западный', 'восточный', 'северо-западный',
         'северо-восточный', 'юго-западный', 'юго-восточный']

d1 = date(year, 1, 1)  # начальная дата
d2 = date(year, 12, 31)  # конечная дата
delta = d2 - d1  # timedelta

p = document.add_paragraph('')

for i in range(delta.days + 1):
    day = d1 + timedelta(i)  # дата
    t = random.randrange(-20, 31)  # температура воздуха
    wind = random.choice(winds)  # направление ветра
    bar = random.randrange(730, 790)  # атмосферно давление
    gygro = random.randrange(70, 100) # влажность воздуха

    day_weather = 'В день {} была температура воздуха {} градусов по Цельсию, {} ветер, ' \
                  'атмосферное давление {} мм рт.ст., влажность воздуха {}%. '.format(day, t, wind, bar, gygro)

    p.add_run(day_weather)

document.save('weather.docx')
